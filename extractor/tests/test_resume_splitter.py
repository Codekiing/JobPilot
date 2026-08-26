from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from resume_splitter import ResumeParser


class ResumeParserTests(unittest.TestCase):
    def setUp(self) -> None:
        self.parser = ResumeParser()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_markdown_resume_is_split_and_saved(self) -> None:
        resume = self.root / "resume.md"
        resume.write_text(
            """# 张三

邮箱：zhangsan@example.com | 电话：13800138000
求职意向：后端工程师

## 教育背景
2020.09-2024.06 示例大学 计算机科学

## 项目经历
2023.01-2023.06 搜索系统
- 将检索延迟降低 30%

## 核心技能
编程语言：Python、Go
数据库：PostgreSQL
""",
            encoding="utf-8",
        )
        document, output = self.parser.parse_and_save(resume, self.root / "out")

        self.assertEqual(document.profile.name, "张三")
        self.assertEqual(document.profile.email, "zhangsan@example.com")
        self.assertEqual(document.profile.phone, "13800138000")
        self.assertEqual(
            [section.type for section in document.sections],
            ["basic_info", "education", "project", "skills"],
        )
        self.assertEqual(document.sections[1].items[0].date, "2020.09-2024.06")
        saved = json.loads((output / "resume.json").read_text(encoding="utf-8"))
        self.assertEqual(saved["schema_version"], "1.0")
        self.assertTrue((output / "sections" / "02-教育背景.md").exists())

    def test_latex_sections_are_preserved(self) -> None:
        resume = self.root / "resume.tex"
        resume.write_text(
            r"""\documentclass{article}
\begin{document}
李四\\
lisi@example.com\\
\section{Education}
2021.09--2025.06 Demo University
\section{Skills}
\textbf{Languages}: Python, Rust
\end{document}
""",
            encoding="utf-8",
        )
        document = self.parser.parse_file(resume)
        self.assertEqual(document.profile.name, "李四")
        self.assertIn("education", [section.type for section in document.sections])
        self.assertIn("skills", [section.type for section in document.sections])

    def test_docx_headings_and_tables_are_extracted(self) -> None:
        from docx import Document

        resume = self.root / "resume.docx"
        source = Document()
        source.add_paragraph("王五")
        source.add_paragraph("wangwu@example.com")
        source.add_heading("工作经历", level=1)
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "2022.01-2024.12"
        table.cell(0, 1).text = "示例科技有限公司 算法工程师"
        source.add_heading("Skills", level=1)
        source.add_paragraph("编程语言：Python")
        source.save(resume)

        document = self.parser.parse_file(resume)
        self.assertEqual(document.profile.name, "王五")
        self.assertEqual([section.type for section in document.sections], ["basic_info", "experience", "skills"])
        self.assertEqual(document.sections[1].items[0].date, "2022.01-2024.12")

    def test_unknown_extension_is_rejected(self) -> None:
        resume = self.root / "resume.xyz"
        resume.write_text("hello", encoding="utf-8")
        with self.assertRaisesRegex(Exception, "不支持的文件格式"):
            self.parser.parse_file(resume)

if __name__ == "__main__":
    unittest.main()
